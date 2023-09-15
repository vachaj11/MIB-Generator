"""This module enables generation of a ``.docx`` file which documents in a human-readable way what was put in MIB tables for a
given set of Tm and Tc packets.

The only (important) method in this module builds up the ``.docx`` file using the not so powerful ``python-docx`` library.
"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches

import mib_generator.parsing.load as load


def gen_doc(Tm_packets=[], Tc_packets=[]):
    """Create a ``.docx`` document which sketches out what entries are in each of the Tm and Tc packets in the passed
    list.

    This method goes through all passed Tm and Tc packets and for each adds an entry to the created ``.docx`` document
    including basic info about the packet and a table of entries within it. The format of the output tries to adhere to a
    standard scheme of an ICD document used normally for these purposes.

    Args:
        Tm_packets (list): List of Tm-packets, each of type :obj:`mib_generator.construction.TM_packet.TM_packet`, from which
            the document is to be generated.
        Tc_packets (list): List of Tc-packets, each of type :obj:`mib_generator.construction.TC_packet.TC_packet`, from which
            the document is to be generated.

    """
    document = Document()
    if Tm_packets:
        document.add_heading("Telemetry packets", 1)
    for i in Tm_packets:
        h = document.add_heading("", 2)
        try:
            t = str(i.pid["PID_TYPE"])
            st = str(i.pid["PID_STYPE"])
            if int(i.pid["PID_PI1_VAL"]) > 0:
                si = ",SID=" + str(i.pid["PID_PI1_VAL"])
            else:
                si = ""
            desc = i.pid["PID_DESCR"]
        except:
            t = "?"
            st = "?"
            si = ""
            desc = "____"
        h.add_run("TM[" + t + "," + st + si + "] " + desc).italic = True
        try:
            document.add_paragraph(
                "APID: " + str(i.pid["PID_APID"]), style="List Bullet"
            )
        except:
            document.add_paragraph("APID: ?", style="List Bullet")
        try:
            document.add_paragraph(
                "MIB name: " + i.tpcf["TPCF_NAME"], style="List Bullet"
            )
        except:
            document.add_paragraph("MIB name: ?", style="List Bullet")
        try:
            document.add_paragraph(
                "MIB SPID: " + str(i.pid["PID_SPID"]), style="List Bullet"
            )
        except:
            document.add_paragraph("MIB SPID: ?", style="List Bullet")
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.autofit = True
        table.rows[0].cells[0].paragraphs[0].add_run("Byte").bold = True
        table.rows[0].cells[1].paragraphs[0].add_run("ID").bold = True
        table.rows[0].cells[2].paragraphs[0].add_run("Size in bites").bold = True
        table.rows[0].cells[3].paragraphs[0].add_run("Description").bold = True
        table.columns[0].width = Cm(3.2)
        table.columns[1].width = Cm(3.2)
        table.columns[2].width = Cm(3.2)
        table.columns[3].width = Cm(5.9)
        for l in range(4):
            table.rows[0].cells[l]._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="C7D9F1"/>'.format(nsdecls("w")))
            )
            table.rows[0].cells[l].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for l in range(len(i.plf)):
            row = table.add_row().cells
            row[0].text = to_bytes(i.positions[l])
            row[1].text = i.pcf[l]["PCF_NAME"]
            row[2].text = str(i.positions[l + 1] - i.positions[l])
            row[3].text = i.pcf[l]["PCF_DESCR"]
            for k in range(4):
                row[k].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        document.add_paragraph()
    if Tc_packets:
        document.add_heading("Telecommand packets", 1)
    for i in Tc_packets:
        h = document.add_heading("", 2)
        try:
            t = str(i.ccf["CCF_TYPE"])
            st = str(i.ccf["CCF_STYPE"])
            desc = i.ccf["CCF_DESCR"]
        except:
            t = "?"
            st = "?"
            desc = "____"
        h.add_run("TC[" + t + "," + st + "] " + desc).italic = True
        try:
            if str(i.ccf["CCF_APID"]):
                document.add_paragraph(
                    "APID: " + str(i.ccf["CCF_APID"]), style="List Bullet"
                )
            else:
                document.add_paragraph("APID: ?", style="List Bullet")
        except:
            document.add_paragraph("APID: ?", style="List Bullet")
        try:
            document.add_paragraph(
                "MIB name: " + str(i.ccf["CCF_CNAME"]), style="List Bullet"
            )
        except:
            document.add_paragraph("MIB name: ?", style="List Bullet")
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.autofit = True
        table.rows[0].cells[0].paragraphs[0].add_run("Byte").bold = True
        table.rows[0].cells[1].paragraphs[0].add_run("ID").bold = True
        table.rows[0].cells[2].paragraphs[0].add_run("Size in bites").bold = True
        table.rows[0].cells[3].paragraphs[0].add_run("Type").bold = True
        table.rows[0].cells[4].paragraphs[0].add_run("Description").bold = True
        table.columns[0].width = Cm(2.7)
        table.columns[1].width = Cm(2.7)
        table.columns[2].width = Cm(2.7)
        table.columns[3].width = Cm(2.7)
        table.columns[4].width = Cm(4.8)
        for l in range(5):
            table.rows[0].cells[l]._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="C7D9F1"/>'.format(nsdecls("w")))
            )
            table.rows[0].cells[l].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for l in range(len(i.cdf)):
            row = table.add_row().cells
            row[0].text = to_bytes(i.positions[l])
            row[1].text = i.cdf[l]["CDF_PNAME"]
            if not row[1].text:
                row[1].text = i.entries[l].name
            row[2].text = str(i.positions[l + 1] - i.positions[l])
            row[3].text = str(i.entries[l].type)
            row[4].text = i.cdf[l]["CDF_DESCR"]
            for k in range(4):
                row[k].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        document.add_paragraph()
    return document


def to_bytes(bits):
    """Give the number of bits in bytes and in some meaning full format.

    Calculates how many bytes are in the passed bits, how many residual bits, and puts together this information into a
    meaningful-looking string.

    Args:
        bits (int): Numeber of bits to be "translated".

    Returns:
        str: The number of bits in bytes and residual bits in understandable form.

    """
    byt = int(bits / 8)
    bit = bits - byt * 8
    if bit == 0:
        return str(byt)
    elif bit == 1:
        return str(byt) + " (+ " + str(bit) + " bit)"
    else:
        return str(byt) + " (+ " + str(bit) + " bits)"
